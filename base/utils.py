# # utils.py
# import os
# import fitz
# from docx import Document
# import spacy



# nlp = spacy.load("en_core_web_sm")
# KNOWN_SKILLS = {
#     "python", "java", "c++", "javascript", "typescript", "ruby", "go", "kotlin", "swift",
#     "react", "angular", "vue.js", "node.js", "django", "flask", "spring boot", "asp.net",
#     "git", "github", "bitbucket", "jenkins", "docker", "kubernetes", "ci/cd", "unit testing",
#     "rest api", "graphql", "mysql", "postgresql", "mongodb", "redis", "sqlite",
#     "fpga", "vhdl", "verilog", "pcb design", "circuit design", "schematic capture",
#     "microcontrollers", "embedded systems", "altium designer", "orcad", "cadence",
#     "signal integrity", "thermal analysis", "emc", "debugging", "hardware prototyping",
#     "windows server", "linux", "unix", "active directory", "vmware", "dns", "dhcp", "vpn",
#     "tcp/ip", "bash", "powershell", "system administration", "itil", "servicenow",
#     "network security", "penetration testing", "hipaa", "gdpr", "azure", "aws", "gcp",
#     "r", "sql", "sas", "scikit-learn", "tensorflow", "pytorch", "machine learning",
#     "deep learning", "regression analysis", "classification", "clustering", "data cleaning",
#     "feature engineering", "nlp", "text mining", "big data", "spark", "hadoop",
#     "data visualization", "tableau", "power bi", "a/b testing", "predictive modeling",
#     "catia", "autocad", "solidworks", "creo", "vehicle dynamics", "powertrain", "adas",
#     "bms", "thermal management", "simulink", "can bus", "diagnostic tools", "fmea", "dfmea",
#     "ppap", "lean manufacturing", "six sigma",
#     "civil 3d", "revit", "staad.pro", "sap2000", "etabs", "hec-ras", "arcgis",
#     "structural analysis", "project scheduling", "cost estimation", "contract administration",
#     "urban planning", "environmental engineering", "site development", "land surveying",
#     "erosion control", "bim", "construction management"
# }


# def extract_text(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == '.pdf':
#         return extract_text_from_pdf(file_path)
#     elif ext == '.docx':
#         return extract_text_from_docx(file_path)
#     else:
#         raise ValueError("Unsupported file format")

# def extract_text_from_pdf(path):
#     text = ""
#     doc = fitz.open(path)
#     for page in doc:
#         text += page.get_text()
#     return text

# def extract_text_from_docx(path):
#     doc = Document(path)
#     return "\n".join([para.text for para in doc.paragraphs])




# # def extract_keywords(text):
# #     doc = nlp(text)
# #     keywords = set()

# #     '''# 1. Add noun_chunks (multi-word phrases)
# #     for chunk in doc.noun_chunks:
# #         phrase = chunk.text.lower().strip()
# #         if len(phrase.split()) > 1:
# #             keywords.add(phrase)'''

# #     # 2. Add meaningful single words
# #     for token in doc:
# #         if token.pos_ in ["NOUN", "PROPN", "VERB"] and token.is_alpha and len(token) > 2:
# #             keywords.add(token.lemma_.lower())

# #     # 3. Match known skill keywords directly in the raw text
# #     text_lower = text.lower()
# #     for skill in KNOWN_SKILLS:
# #         if skill in text_lower:
# #             keywords.add(skill)

# #     return keywords

# # from nltk.tokenize import word_tokenize
# # from nltk.corpus import stopwords
# # import string

# # def extract_keywords(text):
# #     import nltk
# #     nltk.download('punkt') 
# #     # Normalize
# #     KNOWN_SKILLS = {"python", "java", "c++", "catia", "aws", "rest api", "ci/cd", "sql", "matlab", "tensorflow"}  # sample

# #     text = text.lower()

# #     # Tokenize
# #     tokens = word_tokenize(text)

# #     # Remove stopwords and punctuation
# #     stop_words = set(stopwords.words('english'))
# #     filtered_tokens = [
# #         token for token in tokens 
# #         if token not in stop_words and token not in string.punctuation
# #     ]

# #     # Match with known skills
# #     matched_keywords = {skill for skill in KNOWN_SKILLS if skill in " ".join(filtered_tokens)}

# #     return matched_keywords




# # def match_resume_to_job(resume_text, job_desc_text):
# #     resume_kw = extract_keywords(resume_text)
# #     jd_kw = extract_keywords(job_desc_text)
# #     matched = resume_kw & jd_kw
# #     score = len(matched) / len(jd_kw) * 100 if jd_kw else 0
# #     return round(score, 2), matched

# from nltk.tokenize import word_tokenize
# from nltk.corpus import stopwords
# import string
# import nltk

# # Download once at the top-level so it's not inside the function
# # nltk.download('punkt')
# # nltk.download('stopwords')

# # def extract_keywords(text):
# #     # Normalize
# #     text = text.lower()

# #     # Tokenize
# #     tokens = word_tokenize(text)

# #     # Remove stopwords and punctuation
# #     stop_words = set(stopwords.words('english'))
# #     filtered_tokens = [
# #         token for token in tokens
# #         if token not in stop_words and token not in string.punctuation
# #     ]

# #     # Return as a set for matching
# #     return set(filtered_tokens)

# from nltk.tokenize import word_tokenize
# from nltk.corpus import stopwords
# from nltk import pos_tag
# import string
# import nltk

# # Download once
# # nltk.download('punkt')
# # nltk.download('stopwords')
# #nltk.download('averaged_perceptron_tagger')

# import spacy
# from nltk.corpus import stopwords
# import string

# nlp = spacy.load("en_core_web_sm")

# def extract_keywords(text):
#     # Process text with spaCy
#     doc = nlp(text.lower())

#     # Define stopwords and allowed POS types
#     stop_words = set(stopwords.words('english'))
#     allowed_pos = {'NOUN', 'PROPN'}  # Only nouns and proper nouns

#     # Extract keywords
#     keywords = {
#         token.lemma_ for token in doc
#         if token.pos_ in allowed_pos and
#            token.lemma_ not in stop_words and
#            token.lemma_ not in string.punctuation and
#            len(token.lemma_) > 2

#     }

#     return keywords



# def match_resume_to_job(resume_text, job_desc_text):
#     resume_kw = extract_keywords(resume_text)
#     jd_kw = extract_keywords(job_desc_text)
#     matched = resume_kw & jd_kw
#     score = len(matched)
#     return round(score, 2), matched


import os
import fitz  # PyMuPDF
from docx import Document
import spacy
import nltk
from nltk.corpus import stopwords
import string

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Ensure necessary NLTK resources are downloaded
# nltk.download("punkt")
# nltk.download("stopwords")

# Stopwords and punctuation
stop_words = set(stopwords.words("english"))
punctuation = set(string.punctuation)

# === Text Extraction ===
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")

def extract_text_from_pdf(path):
    text = ""
    doc = fitz.open(path)
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(path):
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

# === Hybrid Keyword Extractor ===
def extract_keywords(text):
    doc = nlp(text)
    keywords = set()

    for token in doc:
        if (
            token.is_alpha and                    # ignore numbers/symbols
            token.lemma_ not in stop_words and    # remove stopwords
            token.lemma_ not in punctuation and
            len(token.lemma_) > 2 and             # filter very short words
            token.pos_ in ["PROPN"]       # retain only informative nouns
        ):
            keywords.add(token.lemma_)
    
    return keywords

# === Keyword Matcher ===
def match_resume_to_job(resume_text, job_desc_text):
    resume_kw = extract_keywords(resume_text)
    jd_kw = extract_keywords(job_desc_text)
    matched = resume_kw & jd_kw
    score = len(matched) / len(jd_kw) * 100 if jd_kw else 0
    return round(score, 2), matched
